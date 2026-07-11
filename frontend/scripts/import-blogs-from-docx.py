#!/usr/bin/env python3
"""Import the first 22 ready blogs from 1-22 Blogs.docx + blogs-export.json."""
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
DOCX = Path("/Users/nikhilanand/Downloads/ITR-Startup related/Blogs/1-22 Blogs.docx")
EXPORT = ROOT / "backend/scripts/blogs-export.json"
OUT_JSON = ROOT / "frontend/data/blogs.json"
OUT_IMG_DIR = ROOT / "frontend/public/blogs"

ARTICLE_RE = re.compile(r"^(\d+)\.\s+(.+)$")
SUBHEADING_SKIP = (
    "Encashment During Employment",
    "Encashment at Retirement",
    "Government Employees",
    "Private Employees",
    "Standard Severance",
    "Voluntary Retirement Scheme",
)


def normalize_title(title: str) -> str:
    return re.sub(r"\s+", " ", title.strip().lower())


def para_image_blob(paragraph):
    for blip in paragraph._element.xpath(".//a:blip"):
        embed = blip.get(qn("r:embed"))
        if not embed:
            continue
        part = paragraph.part.related_parts.get(embed)
        if part is None:
            continue
        ext = part.content_type.split("/")[-1]
        if ext == "jpeg":
            ext = "jpg"
        return part.blob, ext
    return None, None


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Missing docx: {DOCX}")

    articles = json.loads(EXPORT.read_text(encoding="utf-8"))[:22]
    title_to_slug = {normalize_title(a["title"]): a["slug"] for a in articles}

    doc = Document(str(DOCX))
    article_paras: list[tuple[int, int, str]] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t or not p.style or p.style.name != "Heading 2":
            continue
        m = ARTICLE_RE.match(t)
        if not m:
            continue
        num = int(m.group(1))
        title = m.group(2).strip()
        if num > 22:
            continue
        if any(skip in title for skip in SUBHEADING_SKIP):
            continue
        article_paras.append((i, num, title))

    # De-dupe by article number (keep first occurrence)
    by_num: dict[int, tuple[int, str]] = {}
    for idx, num, title in article_paras:
        if num not in by_num:
            by_num[num] = (idx, title)
    ordered = [by_num[n] for n in sorted(by_num)]

    OUT_IMG_DIR.mkdir(parents=True, exist_ok=True)
    posts = []

    for idx, (start_idx, title) in enumerate(ordered):
        n = idx + 1
        slug = title_to_slug.get(normalize_title(title))
        if not slug:
            raise SystemExit(f"No slug match for article {n}: {title}")

        src = next(a for a in articles if a["slug"] == slug)
        end_idx = ordered[idx + 1][0] if idx + 1 < len(ordered) else len(doc.paragraphs)

        cover_path = None
        for j in range(start_idx, min(start_idx + 40, end_idx)):
            blob, ext = para_image_blob(doc.paragraphs[j])
            if blob:
                filename = f"{slug}.{ext}"
                (OUT_IMG_DIR / filename).write_bytes(blob)
                cover_path = f"/blogs/{filename}"
                break

        posts.append(
            {
                "slug": slug,
                "title": src["title"],
                "excerpt": src["description"],
                "body": src["body"],
                "tags": src.get("tags") or [],
                "publishedAt": src.get("publishedAt") or "2026-06-15",
                "readMinutes": src.get("readMinutes") or 6,
                "coverImage": cover_path,
                "source": "upload",
            }
        )

    OUT_JSON.write_text(json.dumps(posts, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Wrote {len(posts)} posts to {OUT_JSON}")
    with_images = sum(1 for p in posts if p.get("coverImage"))
    print(f"Cover images: {with_images}/{len(posts)} in {OUT_IMG_DIR}")


if __name__ == "__main__":
    main()
