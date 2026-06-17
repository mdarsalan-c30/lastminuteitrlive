#!/usr/bin/env python3
"""Build a Word document from the exported blog JSON."""
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import os

with open("scripts/blogs-export.json", "r", encoding="utf-8") as f:
    articles = json.load(f)

# Merge per-domain image prompt maps (slug -> {prompt, alt}) if present.
IMAGE_PROMPTS = {}
for dom in ("salaried", "investor", "deductions", "reconciliation"):
    path = f"docs/blog-program/IMAGE_PROMPTS_{dom}.json"
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as jf:
                IMAGE_PROMPTS.update(json.load(jf))
        except (ValueError, OSError):
            pass

# Domain grouping follows the phase6 aggregator order: 30 each.
DOMAINS = [
    ("Salaried & Regime", 0, 30),
    ("Investors & Freelancers", 30, 60),
    ("Deductions & Senior Citizens", 60, 90),
    ("Reconciliation, Notices & Deadlines", 90, 120),
]

ACCENT = RGBColor(0x0B, 0x5C, 0x46)   # deep green
MUTED = RGBColor(0x6B, 0x72, 0x80)    # grey

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline(paragraph, text):
    """Render inline markdown (bold, italic, code, links) into runs."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = ACCENT
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[(.+?)\]\((.+?)\)", part)
            if m:
                label, url = m.group(1), m.group(2)
                r = paragraph.add_run(label)
                r.font.color.rgb = ACCENT
                r.underline = True
                # keep the URL visible for internal links
                if url.startswith("/"):
                    note = paragraph.add_run(f" ({url})")
                    note.font.size = Pt(8)
                    note.font.color.rgb = MUTED
            else:
                paragraph.add_run(part)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def flush_table(rows):
    cells = [
        [c.strip() for c in row.strip().strip("|").split("|")]
        for row in rows
    ]
    # drop the separator row (---|---)
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(cells):
        tr = table.add_row().cells
        for j in range(ncols):
            txt = row[j] if j < len(row) else ""
            para = tr[j].paragraphs[0]
            if i == 0:
                run = para.add_run(txt)
                run.bold = True
            else:
                add_inline(para, txt)
    doc.add_paragraph()


def render_body(body):
    lines = body.split("\n")
    table_buf = []
    for line in lines:
        stripped = line.strip()
        # table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(stripped)
            continue
        elif table_buf:
            flush_table(table_buf)
            table_buf = []

        if not stripped:
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=2)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(stripped.lstrip("> ").strip())
            r.italic = True
            r.font.color.rgb = MUTED
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
    if table_buf:
        flush_table(table_buf)


# ---------- Cover page ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("LastMinute ITR")
run.font.size = Pt(34)
run.font.bold = True
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("SEO Blog Library — 120 Articles for Indian Taxpayers")
r.font.size = Pt(16)
r.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Conversational, jargon-free guides across Salary, Investments, Deductions, "
    "and Reconciliation.\nCompanion content — readers always file on incometax.gov.in."
)
r.font.size = Pt(10)
r.font.color.rgb = MUTED

doc.add_page_break()

# ---------- Table of contents ----------
doc.add_heading("Contents", level=1)
idx = 1
for domain, start, end in DOMAINS:
    h = doc.add_paragraph()
    rr = h.add_run(f"{domain}  ({end - start} articles)")
    rr.bold = True
    rr.font.color.rgb = ACCENT
    for a in articles[start:end]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(a["title"])
        idx += 1

doc.add_page_break()

# ---------- Articles ----------
counter = 1
for domain, start, end in DOMAINS:
    doc.add_heading(domain, level=1)
    intro = doc.add_paragraph()
    ir = intro.add_run(f"{end - start} guides in this section.")
    ir.italic = True
    ir.font.color.rgb = MUTED
    doc.add_paragraph()

    for a in articles[start:end]:
        doc.add_heading(f"{counter}. {a['title']}", level=2)

        m = doc.add_paragraph()
        mr = m.add_run(
            f"Slug: /learn/{a['slug']}   |   "
            f"Cluster: {a['cluster']}   |   "
            f"{a['readMinutes']} min read   |   "
            f"Tags: {', '.join(a['tags'])}"
        )
        mr.font.size = Pt(8.5)
        mr.font.color.rgb = MUTED

        d = doc.add_paragraph()
        dr = d.add_run(a["description"])
        dr.italic = True

        img = IMAGE_PROMPTS.get(a["slug"])
        if img and isinstance(img, dict) and img.get("prompt"):
            ip = doc.add_paragraph()
            lbl = ip.add_run("Image prompt (generate with AI): ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = ACCENT
            pr = ip.add_run(img["prompt"])
            pr.font.size = Pt(9)
            pr.font.color.rgb = MUTED
            if img.get("alt"):
                ap = doc.add_paragraph()
                ar = ap.add_run(f"Alt text: {img['alt']}")
                ar.font.size = Pt(8.5)
                ar.italic = True
                ar.font.color.rgb = MUTED

        render_body(a["body"])

        if a["faqs"]:
            doc.add_heading("Frequently asked questions", level=3)
            for faq in a["faqs"]:
                qp = doc.add_paragraph()
                add_inline(qp, f"**{faq['question']}**")
                ap = doc.add_paragraph()
                add_inline(ap, faq["answer"])

        doc.add_paragraph("\u2014 \u2014 \u2014")
        counter += 1

out = "LastMinute_ITR_Blog_Library_120.docx"
doc.save(out)
print(f"Saved {out} with {counter - 1} articles")
